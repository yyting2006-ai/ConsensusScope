from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_BILINGUAL = ROOT / "docs" / "ConsensusScope_EMNLP_demo_script_2min30_bilingual.docx"
OUT_EN = ROOT / "docs" / "ConsensusScope_EMNLP_demo_script_2min30_en.docx"


SEGMENTS = [
    {
        "time": "0:00-0:20",
        "page": "Problem / Page 1: Review Workspace",
        "operation": [
            "打开 Page 1: Review Workspace。",
            "鼠标停在标题和顶部指标附近。",
            "指一下主流程：single essay, batch review, comparison, queue, evaluation, reports。",
        ],
        "narration": (
            "AI writing feedback can be fluent but unsafe. A model may fix a local grammar issue while also changing a student's intended meaning, "
            "adding unsupported content, or overcorrecting a reasonable ESL draft."
        ),
    },
    {
        "time": "0:20-0:50",
        "page": "Page 2: Single Essay Review",
        "operation": [
            "点击 Page 2: Single Essay Review。",
            "选择内置 demo essay。",
            "展示 assignment prompt、essay text、reviewer settings。",
            "如果页面已有 routed feedback，就直接展示；如果没有，点击生成/路由按钮。",
            "指向 auto-accepted items、teacher-review items、risk score、evidence signal、review explanation。",
        ],
        "narration": (
            "In the single essay window, a teacher can paste an ESL draft, provide the assignment prompt, and generate AI-style feedback candidates. "
            "The system then routes each feedback item before it reaches the student. Each item receives a risk score, evidence signal, review priority, and short explanation for the teacher."
        ),
    },
    {
        "time": "0:50-1:10",
        "page": "Page 3: Batch Review",
        "operation": [
            "点击 Page 3: Batch Review。",
            "展示 packaged synthetic CSV / sample data。",
            "展示 batch summary table 和 routed feedback export。",
            "不需要上传真实文件。",
        ],
        "narration": (
            "The batch window supports the practical classroom workflow: multiple essays can be processed from a CSV, then exported as routed feedback for teacher triage."
        ),
    },
    {
        "time": "1:10-1:30",
        "page": "Page 4: AI Feedback Comparison",
        "operation": [
            "点击 Page 4: AI Feedback Comparison。",
            "指向按 target span 和 issue type 对齐的反馈。",
            "指向 risk level / consensus state / review routing 相关列。",
        ],
        "narration": (
            "The comparison page makes model disagreement visible. Feedback is grouped by target span and issue type, with reviewers, suggestions, risk levels, and consensus state shown together."
        ),
    },
    {
        "time": "1:30-2:00",
        "page": "Page 5: Teacher Queue",
        "operation": [
            "点击 Page 5: Teacher Queue。",
            "停在一个 high-risk item 上，例如 meaning change 或 unsupported claim。",
            "指向 Feedback Safety Graph path、review confidence、evidence signal、priority、explanation。",
            "如果页面有 action 控件，可以选择或展示一个 teacher action。",
        ],
        "narration": (
            "The teacher queue prioritizes high-risk feedback first. Here are four cases: a safe local phrase edit can be accepted; "
            "a thesis-reversing suggestion is routed to review; an unsupported exam-score claim is blocked; and a teacher-dependent punctuation suggestion is now reviewable after our two-teacher diagnostic pilot."
        ),
    },
    {
        "time": "2:00-2:20",
        "page": "Page 6: Effectiveness Evaluation",
        "operation": [
            "点击 Page 6: Effectiveness Evaluation。",
            "指向 action accuracy、risk accuracy、high-risk recall、review recall、auto-accept precision。",
            "如页面有公开学习者语料 benchmark 表，指向 auto share、review share、errors reviewed。",
            "点击 Page 7: Reports，展示 report preview 和 export buttons。",
        ],
        "narration": (
            "The evaluation page separates two kinds of evidence. The synthetic checks verify implementation behavior, while the public learner-corpus benchmark evaluates routing on JFLEG, "
            "CoNLL-2014, FCE, and W&I plus LOCNESS correction data. We also ran a small two-teacher blind Likert pilot over 30 feedback items. "
            "After adding deploy-time signals for teacher-dependent wording, semantic drift, and wrong local corrections, review-needed and unsafe-item recall both reach 1.000. "
            "These results validate graph-backed review routing, not classroom learning outcomes."
        ),
    },
    {
        "time": "2:20-2:30",
        "page": "Page 7: Reports",
        "operation": [
            "停在 Reports 或返回 Page 1。",
            "鼠标不要动，留 1 秒安静收尾。",
        ],
        "narration": "ConsensusScope turns AI feedback into a reviewable teaching workflow: generate, compare, route, review, and export.",
    },
]


def set_run_font(run, font_name: str = "Arial", size: int | None = None, bold: bool | None = None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, font_name: str, size: int, color: str | None = None, bold: bool | None = None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=90, bottom=90, start=120, end=120)


def add_labeled_para(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(label)
    set_run_font(r1, size=10, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_font(r, size=9)


def add_timeline_table(doc: Document):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    headers = ["Time", "Page", "Chinese screen operation", "English narration focus"]
    widths = [0.7, 1.55, 3.15, 1.1]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "E8EEF5")
    for segment in SEGMENTS:
        row = table.add_row()
        values = [
            segment["time"],
            segment["page"],
            "\n".join(segment["operation"]),
            segment["narration"].split(".")[0] + ".",
        ]
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            cell.width = Inches(widths[idx])
            set_cell_text(cell, text, size=8 if idx == 2 else 9)
    doc.add_paragraph()


def add_segment(doc: Document, segment: dict[str, object]):
    h = doc.add_paragraph(style="Heading 2")
    r = h.add_run(f"{segment['time']}  {segment['page']}")
    set_run_font(r, size=12, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("屏幕操作（中文，录屏时照做）")
    set_run_font(r, size=10, bold=True)
    for action in segment["operation"]:
        add_bullet(doc, str(action))

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    r1 = p.add_run("English narration: ")
    set_run_font(r1, size=10, bold=True)
    r2 = p.add_run(str(segment["narration"]))
    set_run_font(r2, size=10)


def build_doc(out_path: Path, bilingual: bool = True) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    set_style_font(styles["Normal"], "Arial", 10, "000000")
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    set_style_font(styles["Heading 1"], "Arial", 15, "2E74B5", True)
    set_style_font(styles["Heading 2"], "Arial", 12, "2E74B5", True)
    set_style_font(styles["Heading 3"], "Arial", 11, "1F4D78", True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("ConsensusScope EMNLP Demo Recording Script")
    set_run_font(r, size=20, bold=True)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_text = "Chinese screen actions + English narration · 2 minutes 30 seconds" if bilingual else "English narration · 2 minutes 30 seconds"
    r = subtitle.add_run(subtitle_text)
    set_run_font(r, size=11)
    r.font.color.rgb = RGBColor.from_string("555555")

    add_labeled_para(doc, "Recording URL: ", "https://demo.consensusscope.cn/")
    add_labeled_para(doc, "Recording target: ", "International conference submission. Use English narration. Keep the final video at or below 2 minutes 30 seconds.")
    add_labeled_para(doc, "Main message: ", "ConsensusScope reviews AI-generated ESL writing feedback before students see it; it is a teacher-in-the-loop routing tool, not an automatic essay scorer or teacher replacement.")

    doc.add_heading("1. One-Page Timeline", level=1)
    add_timeline_table(doc)

    doc.add_heading("2. Read-Aloud Script", level=1)
    for segment in SEGMENTS:
        add_segment(doc, segment)

    doc.add_heading("3. Presenter Safety Checklist", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(
        "Use English narration; show only synthetic demo records; do not reveal API keys or service credentials; "
        "do not claim classroom learning gains; do not frame ConsensusScope as a teacher replacement or automatic essay scorer; "
        "do not bring back the learned meta-judge claim."
    )
    set_run_font(r, size=9)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("ConsensusScope EMNLP demo script · 2:30")
    set_run_font(r, size=8)
    r.font.color.rgb = RGBColor.from_string("777777")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


if __name__ == "__main__":
    build_doc(OUT_BILINGUAL, bilingual=True)
    build_doc(OUT_EN, bilingual=False)
    print(OUT_BILINGUAL)
    print(OUT_EN)

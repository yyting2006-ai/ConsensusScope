from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ConsensusScope_demo_narration_script.docx"


def set_run_font(run, font_name: str = "Microsoft YaHei", size: int | None = None, bold: bool | None = None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, font_name: str, size: int, color: str | None = None, bold: bool | None = None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_labeled_para(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(label)
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11)
    return p


def add_action_steps(doc: Document, steps: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("操作动作：")
    set_run_font(r, size=11, bold=True)
    for step in steps:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(step)
        set_run_font(r, size=10)


def add_script_block(
    doc: Document,
    time: str,
    title: str,
    screen: str,
    narration: str,
    note: str | None = None,
    actions: list[str] | None = None,
):
    h = doc.add_paragraph(style="Heading 2")
    r = h.add_run(f"{time}  {title}")
    set_run_font(r, size=13, bold=True)

    add_labeled_para(doc, "画面操作：", screen)
    if actions:
        add_action_steps(doc, actions)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run("旁白讲稿：")
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(narration)
    set_run_font(r2, size=11)

    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)

    if note:
        add_labeled_para(doc, "录制提示：", note)


def add_small_table(doc: Document, rows: list[tuple[str, str]], widths=(1.55, 4.75)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    for key, value in rows:
        row = table.add_row()
        row.cells[0].width = Inches(widths[0])
        row.cells[1].width = Inches(widths[1])
        set_cell_text(row.cells[0], key, bold=True)
        set_cell_text(row.cells[1], value)
        shade_cell(row.cells[0], "F2F4F7")
    doc.add_paragraph()
    return table


def add_timeline_table(doc: Document, rows: list[tuple[str, str, str, str]]):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    headers = ["时间", "页面", "操作动作", "讲解重点"]
    widths = [0.9, 1.55, 2.25, 1.6]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "E8EEF5")
    for item in rows:
        row = table.add_row()
        for idx, text in enumerate(item):
            cell = row.cells[idx]
            cell.width = Inches(widths[idx])
            set_cell_text(cell, text)
    doc.add_paragraph()
    return table


def build_doc() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], "Microsoft YaHei", 11, "000000")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    set_style_font(styles["Heading 1"], "Microsoft YaHei", 16, "2E74B5", True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Heading 2"], "Microsoft YaHei", 13, "2E74B5", True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], "Microsoft YaHei", 12, "1F4D78", True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("ConsensusScope 项目录屏演示讲稿")
    set_run_font(r, size=22, bold=True)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("面向多大模型协同决策的可靠性评估与动态裁决机制研究")
    set_run_font(r, size=12)
    r.font.color.rgb = RGBColor.from_string("555555")

    add_small_table(
        doc,
        [
            ("适用场景", "项目录屏演示、答辩演示、系统功能展示"),
            ("建议时长", "完整版约 6 分钟；文末附 2 分 30 秒压缩版"),
            ("核心主线", "多模型一致不等于可信；系统关注一致是否可靠、分歧是否有价值、何时需要人工复核。"),
            ("注意事项", "不要把 false consensus 等离线诊断标签说成真实部署时自动知道；它们依赖 gold label，只用于离线分析。"),
        ],
    )

    doc.add_heading("一、录屏前准备", level=1)
    prep = [
        ("启动系统", "运行 start_demo_mac.command 或 streamlit run app/streamlit_app.py，打开本地页面。"),
        ("准备页面", "按侧边栏顺序展示 Home / System Overview、Live Question Mode、Sample Audit Mode、Adjudication Comparison、Risk Dashboard、Model Reliability Dashboard、Case Explorer、Report Export。"),
        ("准备案例", "至少准备一个 false consensus 离线案例和一个 high-risk / high-disagreement 案例。"),
        ("讲解边界", "固定裁判不看 gold answer；动态裁决部署时只看 agreement、diversity、confidence、evidence、minority warning、parse errors 等信号。"),
    ]
    add_small_table(doc, prep, widths=(1.35, 4.95))

    doc.add_heading("二、页面切换与操作总表", level=1)
    add_timeline_table(
        doc,
        [
            ("0:00–0:35", "Page 1", "打开 Home / System Overview；鼠标停在标题与流程图。", "研究问题"),
            ("0:35–1:15", "Page 1", "缓慢下滚，依次指向 API、输出格式、裁决层、看板、导出。", "系统结构"),
            ("1:15–2:00", "Sidebar + Page 2", "展示 API Configuration；切到 Live Question Mode；选择任务并展示输入框。", "API 安全与现场问题"),
            ("2:00–2:45", "Page 2 或 Page 3", "如有 Live 结果就看模型卡片；否则切 Sample Audit，选择 Dataset 和 Sample ID。", "统一输出格式"),
            ("2:45–3:35", "Page 4", "切 Adjudication Comparison；依次指向 majority、fixed judge、dynamic judge。", "三类裁决"),
            ("3:35–4:20", "Page 3 或 Page 7", "打开离线案例；指 gold label、模型答案、risk labels。", "离线诊断标签边界"),
            ("4:20–5:10", "Page 5", "切 Risk Dashboard；停在 low / high risk 结果和 routing 解释。", "实验结果"),
            ("5:10–5:45", "Page 6→7→8", "快速翻三页：模型可靠性、案例浏览、报告导出。", "系统完整性"),
            ("5:45–6:15", "Page 1", "回到 Home，总结后停顿 1 秒结束录屏。", "收束观点"),
        ],
    )

    doc.add_heading("三、完整版录屏讲稿（约 6 分钟）", level=1)

    add_script_block(
        doc,
        "0:00–0:35",
        "开场与研究问题",
        "停留在 Home / System Overview 页面，展示系统标题和整体模块。",
        actions=[
            "打开浏览器中的本地系统页面，确认左侧 Navigation 可见。",
            "在侧边栏点击 Page 1: Home / System Overview。",
            "鼠标停在页面标题 ConsensusScope 和下方系统流程条附近，不要急着滚动。",
            "讲到“多模型一致不等于可信”时，用鼠标轻轻划过流程中的 Adjudication Layer 和 Risk Dashboard。",
        ],
        narration="大家好，我是北京语言大学的 You Tingrui。这个项目叫 ConsensusScope，是一个面向多大模型协同决策的可靠性观测与动态裁决系统。我们关注的问题不是简单地让多个模型投票选出一个答案，而是追问：当多个模型意见一致时，这种一致是否真的可靠；当模型意见分歧时，少数派答案是否可能有价值；以及系统应该在什么时候自动接受、什么时候预警、什么时候交给人工复核。",
        note="开场语速放慢，第一句话就把“多模型一致不等于可信”讲出来。",
    )

    add_script_block(
        doc,
        "0:35–1:15",
        "系统总体设计",
        "继续展示 Home / System Overview，依次指向多模型回答生成、统一输出格式、裁决层、风险仪表盘和报告导出。",
        actions=[
            "仍停留在 Page 1，向下小幅滚动到系统流程或模块说明区域。",
            "鼠标依次指向 API Configuration、Multi-Model Answer Generation、Unified Output Format。",
            "继续指向 Adjudication Layer、Risk Dashboard、Reliability Dashboard、Case Explorer、Report Export。",
            "讲到“四层”时，不要切页，保持在总览页让观众先建立整体结构。",
        ],
        narration="系统整体分为四层。第一层是数据与任务层，支持 TruthfulQA、FEVER、CommonsenseQA，也支持用户在 Live 模式中现场输入问题。第二层是多模型回答生成层，不同模型都输出统一 JSON 格式，包括 answer、reason、confidence 和 evidence。第三层是裁决层，比较 majority vote、fixed judge 和 rule-based dynamic adjudication。第四层是可视化与审计层，用风险仪表盘、案例浏览器和报告导出，把模型的决策过程展示出来。",
    )

    add_script_block(
        doc,
        "1:15–2:00",
        "API Configuration 与 Live Question Mode",
        "切换到 Live Question Mode，展示 API Configuration，说明 Mode A 和 Mode B。",
        actions=[
            "在侧边栏 API Configuration 区域停留，展示 Mode A / Mode B 的选择。",
            "如果是公开录屏，选择或展示 Mode B，但不要输入真实 API key；如果是本地演示，可说明 Mode A 读取本地环境变量。",
            "在侧边栏 Navigation 点击 Page 2: Live Question Mode。",
            "在 Task type 下拉框选择一个任务类型，例如 Factual QA 或 Multiple choice。",
            "在 question 输入框中放入一个短问题；如果不实际调用 API，就只展示输入框和 Run Live Comparison 按钮。",
        ],
        narration="这里是现场问题模式。系统提供两种 API 配置方式。Mode A 适合现场演示，读取本地环境变量或部署 secrets 中已有的 API key；Mode B 适合公开部署，用户临时输入自己的 API key。需要强调的是，API key 不会写进论文，也不会硬编码进代码。用户输入问题后，可以选择多个 OpenAI-compatible 模型同时作答，系统会统一解析它们的答案、理由、置信度和证据，再进入后续裁决流程。",
        note="如果录屏时不想真实调用 API，可以直接说明系统也支持使用保存的 no-API demo 数据。",
    )

    add_script_block(
        doc,
        "2:00–2:45",
        "多模型输出与统一格式",
        "展示 Live Question Mode 或 Sample Audit Mode 中的模型输出卡片。",
        actions=[
            "如果 Live 模式有结果，滚动到模型输出卡片区域，逐个停留在 answer、reason、confidence、evidence 字段。",
            "如果没有现场 API，切到 Page 3: Sample Audit Mode。",
            "在 Dataset 下拉框选择 fever 或 commonsenseqa。",
            "在 Sample ID 下拉框选择一个已有样本，然后滚动到 model outputs 表格或模型卡片。",
            "鼠标依次停在 answer、reason、confidence、evidence 四类字段上，说明统一输出格式。",
        ],
        narration="这一部分展示多模型回答生成模块。每个模型不是只给一个自由文本答案，而是被要求输出结构化字段。answer 表示最终答案，reason 是简短理由，confidence 是模型自报置信度，evidence 是它声称支持答案的依据。统一格式的好处是，后续系统可以横向比较不同模型：哪些模型答案一致，哪些模型有分歧，哪些模型置信度高但证据弱，哪些输出发生了解析错误。",
    )

    add_script_block(
        doc,
        "2:45–3:35",
        "裁决层：Majority Vote、Fixed Judge、Dynamic Judge",
        "切换到 Adjudication Comparison，展示三种裁决方法的对比。",
        actions=[
            "在侧边栏点击 Page 4: Adjudication Comparison。",
            "先把鼠标停在 Majority Vote 区域，讲它是 baseline。",
            "再移动到 Fixed Judge 区域，讲默认 judge provider 和 deepseek-chat。",
            "最后移动到 Dynamic Rule-Based Judge 区域，讲 agreement、diversity、confidence、evidence、parse errors 等风险信号。",
            "如果页面较长，缓慢向下滚动一次，确保三种方法都被录到。",
        ],
        narration="裁决层包含三种主要方法。第一种是 majority vote，也就是多数投票，它适合作为最直观的 baseline，但它的问题是可能产生错误共识。第二种是 fixed judge，当前默认使用 judge provider 下的 deepseek-chat，也可以通过 JUDGE_MODEL 替换。固定裁判看到的是题目、选项，以及其他模型的 answer、reason、confidence 和 evidence；它不看 gold answer，也不看 gold label。第三种是 rule-based dynamic adjudication，它不依赖事后标签，而是在部署时根据 agreement rate、answer diversity、confidence distribution、evidence availability、minority warning 和 parse errors 这些信号，给出风险等级和裁决建议。",
        note="这里要避免说 fixed judge 一定更可信，只说它是可比较的裁判基线。",
    )

    add_script_block(
        doc,
        "3:35–4:20",
        "Sample Audit：离线诊断标签的正确说法",
        "进入 Sample Audit Mode 或 Case Explorer，打开一个 false consensus 案例。",
        actions=[
            "在侧边栏点击 Page 3: Sample Audit Mode。",
            "Dataset 建议选择 fever；Sample ID 选择一个已知离线案例，若默认样本已展示风险标签就保持默认。",
            "先指向 gold answer / gold label，再指向 majority answer 和 dynamic answer。",
            "滚动到 risk labels 或 diagnostic labels 区域，停留 2–3 秒。",
            "讲到“离线诊断标签”时，鼠标不要点 Run，只展示保存结果，避免观众误解为实时知道 gold。",
        ],
        narration="在样本审计页面，我们可以看到一个具体问题下多个模型的输出，以及 majority vote、fixed judge 和 dynamic adjudication 的结果。这里出现的 false consensus、minority correct、true consensus、confidence mismatch 等标签，是离线诊断标签。也就是说，它们依赖数据集中的 gold answer 或 gold label，只能在实验分析和案例复盘时使用。真实部署时，系统不能提前知道一个共识是否一定是 false consensus。因此，系统真正用于部署时依赖的是可观测风险信号，而不是事后标签。",
        note="这一段很重要，可以帮助避免评审质疑“系统怎么提前知道 false consensus”。",
    )

    add_script_block(
        doc,
        "4:20–5:10",
        "Risk Dashboard 与实验结果",
        "切换到 Risk Dashboard，展示风险等级、错误率和 review-routing 结果。",
        actions=[
            "在侧边栏点击 Page 5: Risk Dashboard。",
            "先停在 risk-level 表格或图表，指向 low / medium / high 三个风险等级。",
            "讲到 91.8% 时，把鼠标停在 low-risk accuracy 附近。",
            "讲到 95.2% 时，把鼠标停在 high-risk error rate 附近。",
            "如果页面有 review-routing 或 error-capture 信息，滚动到该区域并停留。",
        ],
        narration="在 1000 个样本、4000 条模型输出的 pilot 实验中，我们没有把 TruthfulQA、FEVER 和 CommonsenseQA 混在一起报告 macro-F1，因为它们的标签空间不同。论文中改用三类评价：第一是按数据集报告 final-answer accuracy；第二是风险分层质量；第三是 review-routing utility。最关键的结果是，rule-based dynamic adjudication 找到的 low-risk 子集准确率达到 91.8%，而 high-risk 子集错误率达到 95.2%。如果把 medium 和 high risk 样本交给人工复核，可以覆盖 83% 的样本，并捕获 97.2% 的最终答案错误。这说明动态裁决的主要价值不是宣称比所有方法都更会选答案，而是把风险集中出来，帮助系统决定何时需要人工复核。",
    )

    add_script_block(
        doc,
        "5:10–5:45",
        "Model Reliability Dashboard、Case Explorer 与 Report Export",
        "依次快速展示 Model Reliability Dashboard、Case Explorer 和 Report Export 页面。",
        actions=[
            "点击 Page 6: Model Reliability Dashboard，停留在模型可靠性统计或图表上。",
            "点击 Page 7: Case Explorer，在 Inspect case 下拉框选择一个案例，展示案例细节。",
            "点击 Page 8: Report Export，展示 Download system_summary.json、Download method_metrics.csv、Download risk_labels.csv 等按钮。",
            "不要真的下载太多文件；最多点击一次下载按钮演示即可。",
        ],
        narration="除了单个样本，系统还提供模型可靠性看板，用来观察不同模型在不同任务上的表现；案例浏览器可以快速筛选高风险样本、错误共识样本和少数派正确样本；报告导出页面可以把当前审计结果导出为 Markdown、JSON 或 CSV，方便后续写论文、做复核或提交 demo 材料。这些功能让系统不只是一个实验脚本，而是一个可以复现实验、检查案例、支持项目汇报的交互式工具。",
    )

    add_script_block(
        doc,
        "5:45–6:15",
        "总结与边界",
        "回到 Home / System Overview 或 Report Export 页面，停留在系统总览。",
        actions=[
            "点击 Page 1: Home / System Overview，回到系统总览。",
            "画面停在标题和系统流程处，鼠标不要再频繁移动。",
            "说完最后一句后停顿 1 秒再结束录屏，避免结尾被截断。",
        ],
        narration="总结来说，ConsensusScope 不是一个新的大模型，也不是一个自动真理判定器。它是一个多模型协同决策的可靠性观测层和风险分流层。它帮助我们看到：模型一致时是否可靠，模型分歧时是否需要保留少数派意见，系统什么时候可以自动接受，什么时候必须预警或交给人工复核。下一步工作可以扩展更多模型、更多数据集，并进一步验证 review-routing 在真实应用场景中的效果。",
    )

    doc.add_heading("四、2 分 30 秒压缩版旁白", level=1)
    compressed = [
        ("0:00–0:20", "大家好，我是北京语言大学的 You Tingrui。ConsensusScope 是一个面向多大模型协同决策的可靠性观测与动态裁决系统。它的核心观点是：多模型一致可以作为证据，但不能直接等同于可信。"),
        ("0:20–0:50", "系统首先统一记录多个模型的结构化输出，包括 answer、reason、confidence 和 evidence。这样我们不仅能看到最后答案，还能比较模型之间是否一致、证据是否充分、置信度是否匹配。"),
        ("0:50–1:25", "裁决层比较三种方法：majority vote、fixed judge 和 rule-based dynamic adjudication。固定裁判默认使用 deepseek-chat，不接收 gold answer，只根据题目和其他模型的输出作出裁决。动态裁决则根据 agreement rate、answer diversity、confidence、evidence 和 parse errors 等部署时可观测信号，给出风险等级。"),
        ("1:25–2:00", "这里要区分两类信息：false consensus、minority correct 等是离线诊断标签，只有在有 gold label 的实验中才能计算；真实部署时，系统依赖的是可观测风险信号。我们的 pilot 实验包含 1000 个样本和 4000 条模型输出，动态裁决识别出的 low-risk 子集准确率为 91.8%，high-risk 子集错误率为 95.2%。"),
        ("2:00–2:30", "因此，ConsensusScope 的价值不是替代人工，也不是声称某个模型一定正确，而是帮助研究者和开发者审计多模型决策过程，判断什么时候可以接受，什么时候应该预警、重查或转人工复核。"),
    ]
    for t, text in compressed:
        add_labeled_para(doc, f"{t}：", text)

    doc.add_heading("五、国际会议语言建议", level=1)
    add_labeled_para(
        doc,
        "结论：",
        "如果目标是 EMNLP / ACL 这类国际会议，正式提交的视频旁白、poster 讲解和现场 demo 建议使用英文。中文可以用于校内项目录屏、内部预演或中文版材料，但不建议作为正式国际会议答辩语言。",
    )
    add_labeled_para(
        doc,
        "原因：",
        "EMNLP System Demonstrations 要求提交带音频旁白的 screencast，并且接受后至少一位作者要在会议 demo session 做 live demo 且配 poster。虽然官方页面通常不会逐字写“禁止中文”，但 ACL/EMNLP 的论文、poster、demo 交流语境默认是英文；用中文会显著影响评审和现场交流。",
    )
    add_labeled_para(
        doc,
        "建议做法：",
        "准备两套版本：中文完整版用于校内项目验收和练习；英文 2.5 分钟版用于国际会议提交。英文版可以照着中文讲稿翻译，不需要复杂修辞，重点是清楚说明 problem、system workflow、risk signals、evaluation 和 limitations。",
    )

    doc.add_heading("六、录屏检查清单", level=1)
    checklist = [
        ("开场是否清楚", "第一分钟内讲出“多模型一致不等于可信”。"),
        ("API 安全是否说明", "说明 Mode A / Mode B，强调 API key 不写论文、不硬编码。"),
        ("Fixed judge 是否透明", "说明默认模型、输入 prompt 内容、不看 gold answer、保存结果作为复现实验产物。"),
        ("指标是否稳妥", "不要主讲混合任务 macro-F1；主讲 dataset accuracy、risk stratification、review-routing utility。"),
        ("标签边界是否清楚", "区分 deploy-time risk signals 与 offline diagnostic labels。"),
        ("结尾是否克制", "强调系统是观测与风险分流工具，不是 truth oracle。"),
    ]
    add_small_table(doc, checklist, widths=(1.6, 4.7))

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("ConsensusScope demo narration script")
    set_run_font(r, size=9)
    r.font.color.rgb = RGBColor.from_string("777777")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
